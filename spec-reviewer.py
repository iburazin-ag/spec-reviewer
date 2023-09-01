import sys, os, subprocess
import argparse
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement

def parse_arguments():
    parser = argparse.ArgumentParser(description="Spec Scanner with Optional Checks")
    parser.add_argument("file_path", help="Path to the Word document file")
    parser.add_argument("--skip-line-breaks", action="store_true", help="Skip line break check")
    parser.add_argument("--skip-formatting", action="store_true", help="Skip formatting check")
    return parser.parse_args()

def is_empty_cell(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            xml_content = run._r.xml
            if xml_content.strip(): 
                return False  
    return True

def comment_formatting(paragraph, finding_text):
    paragraph.add_run().add_break()
    paragraph.add_run(finding_text).bold = True
    paragraph.runs[-1].font.color.rgb = RGBColor(0xFF, 0, 0)

def check_for_existing_findings(cell, finding_text=None):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if run.bold and run.font.color.rgb == RGBColor(0xFF, 0, 0):
                text = run.text.strip()
                if finding_text is None:
                    if text.isupper():
                        run.font.underline = True
                        return True
                elif finding_text in text and text.isupper():
                    run.font.underline = True
                    return True
    return False

def mark_empty_cells(cell):
    paragraph = cell.paragraphs[0]
    if is_empty_cell(cell):
        finding_text = "EMPTY"
        comment_formatting(paragraph, finding_text)
        return True
    return False

def check_and_mark_cdash_cells(cell):
    cdash_text = cell.text.strip()
    paragraph = cell.paragraphs[-1]
   
    if "-" in cdash_text and (" -" in cdash_text or "- " in cdash_text):
        finding_text = "\nREDUNDANT SPACES"
        if not check_for_existing_findings(cell, finding_text.strip()):
            comment_formatting(paragraph, finding_text)
            return True

    if "—" in cdash_text:
        finding_text = "\nDASH INSTEAD OF HYPHEN"
        if not check_for_existing_findings(cell, finding_text.strip()):
            comment_formatting(paragraph, finding_text)
            return True
    
    if "-" not in cdash_text and "N/A" not in cdash_text and "—" not in cdash_text:
        finding_text = "\nMISSING HYPHEN"
        if not check_for_existing_findings(cell, finding_text.strip()):
            comment_formatting(paragraph, finding_text)
            return True
    
    return False

def find_cdash_column(table):
    for col_idx, cell in enumerate(table.rows[0].cells):
        if "CDASH" in cell.text.upper():
            return col_idx
    return None

def check_and_mark_alignment_issue(cell, last_column_cell):
    alignment_missing = False
    for paragraph in cell.paragraphs:
        alignment = paragraph.alignment
        paragraph = last_column_cell.paragraphs[-1]
        finding_text = "\nMISSING ALIGNMENT/FORMATTING COMMENT"
        finding_exists =check_for_existing_findings(last_column_cell, finding_text.strip())

        if not skip_formatting and alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
            if "center aligned" not in last_column_cell.text and not finding_exists:
                comment_formatting(paragraph, finding_text)
                alignment_missing = True
                break
    return alignment_missing

def check_line_breaks(cell):
    finding_text = "\nCHECK LINE BREAKS"
    finding_exists = check_for_existing_findings(cell, finding_text.strip())

    if not skip_line_breaks and '\n' in cell.text and not finding_exists:
        paragraph = cell.paragraphs[-1]
        comment_formatting(paragraph, finding_text)
        return True
    return False
  
if __name__ == "__main__":
    args = parse_arguments()
    file_path = args.file_path
    skip_line_breaks = args.skip_line_breaks
    skip_formatting = args.skip_formatting

    if len(sys.argv) < 2:
        print("Usage: python3 word_table_scanner.py <local_file_path> optional-skip-flag")
    elif not os.path.exists(file_path):
        print("Cannot find the provided file. Please check the file name and path.")
    else:
        document = Document(file_path)
        modified = False
        
        for table in document.tables:
            cdash_col_idx = find_cdash_column(table)

            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    if check_for_existing_findings(cell):
                        modified = True
                        break

                    if is_empty_cell(cell):
                        mark_empty_cells(cell)
                        modified = True
                        break

            for row_idx, row in enumerate(table.rows[1:], start=1):
                alignment_issue_cell = row.cells[-1]
                
                if cdash_col_idx is not None:
                    cdash_cell = row.cells[cdash_col_idx]
                    
                if check_and_mark_cdash_cells(cdash_cell):
                    modified = True
                    break

                for col_idx, cell in enumerate(row.cells[:-1], start=1):  # Exclude the last column
                    alignment_issue = check_and_mark_alignment_issue(cell, alignment_issue_cell)

                    if check_line_breaks(cell):
                        modified = True
                        break

                    if alignment_issue:
                        modified = True
                        break   
        
        if modified:
            document.save(file_path)  
            print("\033[91mFindings recorded in the file.\033[0m")
            try:
                subprocess.run(["open", file_path], check=True)
            except subprocess.CalledProcessError:
                print("Failed to open the modified file.")
        else:
            print("\033[92mNo findings found!\033[0m")
